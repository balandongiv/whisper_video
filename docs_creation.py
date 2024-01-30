import os
from docx import Document
from docx.shared import Inches
from PIL import Image
from support_file import clean_and_format_string


def create_docx_for_each_title(titles, err_or, save_directory):
    """
    Create a DOCX file for each title.
    """
    sanitized_title = clean_and_format_string(titles)
    # gdrive = '/content/drive/MyDrive/'
    # save_directory = gdrive if os.path.exists(gdrive) else local_save_directory
    filename = os.path.join(save_directory, f'{sanitized_title}.docx')

    if not os.path.exists(filename):
        doc = Document()
        doc.add_paragraph(titles)
        # doc.add_paragraph(err_or)  # Uncomment if needed
        doc.save(filename)
        print(f"Document '{filename}' created.")
    else:
        print(f"Document '{filename}' already exists. Skipping.")


def calculate_image_dimensions(image_path, max_width, max_height):
    """
    Calculate new dimensions for an image, maintaining aspect ratio.
    """
    with Image.open(image_path) as img:
        width, height = img.size

    width_ratio = max_width / width
    height_ratio = max_height / height
    scale_factor = min(width_ratio, height_ratio)

    return width * scale_factor, height * scale_factor


def add_images_to_word_document(df, word_fname, max_width, max_height, video_urlxs, num_images=10):
    """
    Add images and transcriptions to a Word document.
    """
    print('Now I am saving all the transcription and images in word.')

    new_width, new_height = calculate_image_dimensions(df.iloc[0]['image_path'], max_width, max_height)
    doc = Document()
    doc.add_paragraph('The video was downloaded from the following url:')
    doc.add_paragraph(video_urlxs)
    doc.add_paragraph("\n")

    for index, row in df.iterrows():
        try:
            image_path = row['image_path']
            doc.add_picture(image_path, width=Inches(new_width), height=Inches(new_height))
        except Exception as e:
            print(f"An error occurred while adding image: {e}")
            continue

        doc.add_paragraph(row['Transcription']['text'])
        doc.add_paragraph("\n")

    doc.save(word_fname)
    print('Completed saving all the transcription and images in word.')
    print(f'The file is saved as {word_fname}')
